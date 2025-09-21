import os, io

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    from docx import Document
except ImportError:
    Document = None


def clean_text(text: str) -> str:
    """Normalize whitespace & lowercase text."""
    if not text:
        return ""
    return " ".join(text.split()).lower()


def parse_jd_file(file_path: str) -> str:
    """Extract text from JD file given its path."""
    if not os.path.exists(file_path):
        return ""

    ext = file_path.rsplit(".", 1)[-1].lower()
    raw_text = ""

    try:
        if ext == "pdf" and fitz:
            doc = fitz.open(file_path)
            raw_text = "\n".join([page.get_text() for page in doc])
        elif ext == "docx" and Document:
            doc = Document(file_path)
            raw_text = "\n".join([p.text for p in doc.paragraphs])
        elif ext == "txt":
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                raw_text = f.read()
        else:
            return ""
    except Exception as e:
        print(f"[jd_parser] failed to parse {file_path}: {str(e).lower()}")
        return ""

    return clean_text(raw_text)


def parse_jd_bytes(file_bytes: bytes, filename: str) -> str:
    """Extract text from JD file given raw bytes + filename."""
    ext = filename.rsplit(".", 1)[-1].lower()
    raw_text = ""

    try:
        if ext == "pdf" and fitz:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            raw_text = "\n".join([page.get_text() for page in doc])
        elif ext == "docx" and Document:
            doc = Document(io.BytesIO(file_bytes))
            raw_text = "\n".join([p.text for p in doc.paragraphs])
        elif ext == "txt":
            raw_text = file_bytes.decode("utf-8", errors="ignore")
        else:
            return ""
    except Exception as e:
        print(f"[jd_parser] failed to parse bytes for {filename}: {str(e).lower()}")
        return ""

    return clean_text(raw_text)


def extract_text_from_jd(file_bytes: bytes = None, filename: str = None, file_path: str = None) -> str:
    """
    Unified extractor for JD text.
    - file_path → parse_jd_file
    - file_bytes + filename → parse_jd_bytes
    """
    if file_path:
        return parse_jd_file(file_path)
    elif file_bytes and filename:
        return parse_jd_bytes(file_bytes, filename)
    return ""
