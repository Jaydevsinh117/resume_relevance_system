import os
import io

# optional dependencies
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    from docx import Document
except ImportError:
    Document = None


# --------------------------
# helper: clean text
# --------------------------
def clean_text(text: str) -> str:
    """Normalize whitespace and lowercase text."""
    if not text:
        return ""
    return " ".join(text.split()).lower()


# --------------------------
# parse Resume file from path
# --------------------------
def parse_resume_file(file_path: str) -> str:
    """
    Extract text from a resume file (.pdf or .docx) by path.
    """
    if not os.path.exists(file_path):
        print(f"[resume_parser] file not found: {file_path}")
        return ""

    ext = file_path.rsplit(".", 1)[-1].lower()
    raw_text = ""

    try:
        if ext == "pdf" and fitz:
            with fitz.open(file_path) as doc:
                raw_text = "\n".join([page.get_text() for page in doc])
        elif ext == "docx" and Document:
            doc = Document(file_path)
            raw_text = "\n".join([p.text for p in doc.paragraphs])
        else:
            print(f"[resume_parser] unsupported file type: {ext}")
            return ""
    except Exception as e:
        print(f"[resume_parser] failed to parse {file_path}: {str(e).lower()}")
        return ""

    return clean_text(raw_text)


# --------------------------
# parse Resume file from bytes
# --------------------------
def parse_resume_bytes(file_bytes: bytes, filename: str) -> str:
    """
    Extract text from Resume file given raw bytes + filename.
    """
    ext = filename.rsplit(".", 1)[-1].lower()
    raw_text = ""

    try:
        if ext == "pdf" and fitz:
            with fitz.open(stream=file_bytes, filetype="pdf") as doc:
                raw_text = "\n".join([page.get_text() for page in doc])
        elif ext == "docx" and Document:
            doc = Document(io.BytesIO(file_bytes))
            raw_text = "\n".join([p.text for p in doc.paragraphs])
        else:
            print(f"[resume_parser] unsupported file type: {ext}")
            return ""
    except Exception as e:
        print(f"[resume_parser] failed to parse bytes for {filename}: {str(e).lower()}")
        return ""

    return clean_text(raw_text)


# --------------------------
# unified entrypoint
# --------------------------
def extract_text_from_resume(
    file_bytes: bytes = None, filename: str = None, file_path: str = None
) -> str:
    """
    Unified extractor for Resume text. Works with either raw bytes + filename or file path.

    - If file_path is given → uses parse_resume_file
    - Else → uses parse_resume_bytes
    """
    if file_path:
        return parse_resume_file(file_path)
    elif file_bytes and filename:
        return parse_resume_bytes(file_bytes, filename)
    else:
        print("[resume_parser] no input provided (file_path or file_bytes required)")
        return ""
